"""Generation tools for documents, sections, charts, and tables."""

import asyncio
import logging
from io import BytesIO
from typing import Any

from openai import AsyncOpenAI

# Optional weasyprint import - allows tests to run without GTK libraries
try:
    from weasyprint import CSS, HTML

    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    # OSError: missing GTK system libraries on Windows
    WEASYPRINT_AVAILABLE = False
    HTML = None  # type: ignore
    CSS = None  # type: ignore

from ..metrics import get_metrics
from ..metrics.llm_cost_calculator import LLMCostCalculator
from ..models.tool_schemas import (
    AssembleDocumentInput,
    AssembleDocumentOutput,
    CreateChartInput,
    CreateChartOutput,
    FormatTableInput,
    FormatTableOutput,
    GenerateSectionInput,
    GenerateSectionOutput,
    ToolCategory,
)
from ..services.llm_budget_guardrails import LLMBudgetExceededError, get_llm_budget_guardrails
from .registry import BaseTool

logger = logging.getLogger(__name__)


class GenerateSectionTool(BaseTool):
    """Generate narrative sections for business cases using LLM."""

    name = "generate_section"
    category = ToolCategory.GENERATION
    description = "Generates narrative sections for business cases using LLM"
    input_schema = GenerateSectionInput
    output_schema = GenerateSectionOutput
    timeout_seconds = 30

    # Section templates with prompts
    SECTION_TEMPLATES = {
        "executive_summary": """Write a 2-3 paragraph executive summary for a business case.
Context: {context}
Tone: {tone}
Include: Current situation, proposed solution, expected ROI, and recommendation.
Maximum {max_length} words.""",
        "current_state": """Describe the current state analysis for a business case.
Context: {context}
Tone: {tone}
Cover: Current pain points, inefficiencies, and business impact.
Maximum {max_length} words.""",
        "proposed_solution": """Describe the proposed solution for a business case.
Context: {context}
Tone: {tone}
Cover: Capabilities offered, implementation approach, and alignment with needs.
Maximum {max_length} words.""",
        "roi_analysis": """Present the ROI analysis for a business case.
Context: {context}
Tone: {tone}
Cover: Investment required, expected returns, payback period, and risk-adjusted NPV.
Maximum {max_length} words.""",
        "implementation": """Outline the implementation plan for a business case.
Context: {context}
Tone: {tone}
Cover: Phases, timeline, key milestones, and resource requirements.
Maximum {max_length} words.""",
        "next_steps": """List recommended next steps for a business case.
Context: {context}
Tone: {tone}
Cover: Immediate actions, decision points, and timeline for approval.
Maximum {max_length} words.""",
    }

    async def _call_llm(self, prompt: str, max_tokens: int = 1000) -> str:
        """Call LLM to generate content."""
        api_key = self.config.get("openai_api_key") if self.config else None
        tenant_id = str(self.config.get("tenant_id", "unknown")) if self.config else "unknown"
        initial_model = "gpt-4o"
        decision = await get_llm_budget_guardrails().precheck_or_raise(tenant_id, initial_model)
        if decision.throttle_seconds > 0:
            await asyncio.sleep(decision.throttle_seconds)

        client = AsyncOpenAI(api_key=api_key)

        response = await client.chat.completions.create(
            model=decision.model,
            messages=[
                {
                    "role": "system",
                    "content": "You are a professional business writer creating business case content.",
                },
                {"role": "user", "content": prompt},
            ],
            temperature=0.3,
            max_tokens=max_tokens,
        )

        prompt_tokens = response.usage.prompt_tokens if response.usage else 0
        completion_tokens = response.usage.completion_tokens if response.usage else 0
        cost = LLMCostCalculator().calculate_cost(
            provider="openai",
            model=decision.model,
            prompt_tokens=prompt_tokens,
            completion_tokens=completion_tokens,
        )
        await get_llm_budget_guardrails().record_usage(tenant_id=tenant_id, cost_usd=cost)

        metrics = get_metrics()
        if metrics:
            metrics.record_llm_cost(
                provider="openai",
                model=decision.model,
                tenant_id=tenant_id,
                cost=cost,
                prompt_tokens=prompt_tokens,
                completion_tokens=completion_tokens,
                status="throttled" if decision.escalation_required else "success",
            )

        if decision.escalation_required:
            logger.warning(
                "Tenant approaching LLM budget soft cap; escalation required",
                extra={"tenant_id": tenant_id, "model": decision.model},
            )

        return response.choices[0].message.content.strip()

    async def execute(self, input_data: GenerateSectionInput) -> GenerateSectionOutput:
        """Generate a business case section."""
        template = self.SECTION_TEMPLATES.get(
            input_data.section_type, self.SECTION_TEMPLATES["executive_summary"]
        )

        # P1-12 FIX: Wrap user context in delimiters to prevent prompt injection
        context_items = [f"{k}: {v}" for k, v in input_data.context.items()]
        context_str = "\n".join(context_items)
        context_delimited = f"<<<USER_CONTEXT>>>{context_str}<<</USER_CONTEXT>>>"

        # P1-12 FIX: Sanitize tone parameter
        allowed_tones = {"professional", "casual", "technical", "executive"}
        safe_tone = input_data.tone if input_data.tone in allowed_tones else "professional"

        prompt = template.format(
            context=context_delimited, tone=safe_tone, max_length=input_data.max_length
        )

        try:
            content = await self._call_llm(prompt, max_tokens=input_data.max_length * 2)
        except LLMBudgetExceededError as e:
            logger.error("LLM budget cap exceeded during section generation: %s", e)
            # CONTRACT_EXCEPTION AP-7: Return structured error, don't raise
            return GenerateSectionOutput(
                content="",
                word_count=0,
                key_points=[],
                error=f"Section generation blocked by tenant budget guardrail: {e}"
            )
        except Exception as e:
            logger.error(f"LLM generation failed: {e}")
            # CONTRACT_EXCEPTION AP-7: Return structured error, don't raise
            return GenerateSectionOutput(
                content="",
                word_count=0,
                key_points=[],
                error=f"Failed to generate section: {e}"
            )

        word_count = len(content.split())

        # Extract key points (simplified implementation)
        key_points = content.split("\n")[:3] if "\n" in content else [content[:100]]
        key_points = [p.strip() for p in key_points if p.strip()]

        return GenerateSectionOutput(
            content=content, word_count=word_count, key_points=key_points[:5]
        )


class CreateChartTool(BaseTool):
    """Create charts and visualizations."""

    name = "create_chart"
    category = ToolCategory.GENERATION
    description = "Creates charts and visualizations for business cases"
    input_schema = CreateChartInput
    output_schema = CreateChartOutput
    timeout_seconds = 20

    async def execute(self, input_data: CreateChartInput) -> CreateChartOutput:
        """Create chart visualization."""
        chart_type = input_data.chart_type
        data = input_data.data
        title = input_data.title

        # Build chart data structure
        chart_data = {
            "type": chart_type,
            "title": title,
            "data": data,
            "config": input_data.config,
        }

        # Add type-specific formatting
        if chart_type == "bar":
            chart_data["x_axis"] = [d.get("label", d.get("category", "")) for d in data]
            chart_data["y_axis"] = [d.get("value", 0) for d in data]

        elif chart_type == "pie":
            chart_data["slices"] = [
                {"label": d.get("label", ""), "value": d.get("value", 0)} for d in data
            ]

        elif chart_type == "line":
            chart_data["series"] = [
                {"name": d.get("series", "Series 1"), "x": d.get("x", []), "y": d.get("y", [])}
                for d in data
            ]

        elif chart_type == "table":
            chart_data["headers"] = input_data.config.get("headers", [])
            chart_data["rows"] = [list(d.values()) for d in data]

        return CreateChartOutput(
            chart_data=chart_data,
            image_data=None,  # Would be generated in production
            image_url=None,
        )


class FormatTableTool(BaseTool):
    """Format data as HTML/Markdown/CSV tables."""

    name = "format_table"
    category = ToolCategory.GENERATION
    description = "Formats data as HTML, Markdown, or CSV tables"
    input_schema = FormatTableInput
    output_schema = FormatTableOutput
    timeout_seconds = 5

    async def execute(self, input_data: FormatTableInput) -> FormatTableOutput:
        """Format table in specified format."""
        headers = input_data.headers
        rows = input_data.rows
        fmt = input_data.format

        # Sort if requested
        if input_data.sort_column is not None and 0 <= input_data.sort_column < len(headers):
            rows = sorted(rows, key=lambda r: r[input_data.sort_column])

        if fmt == "html":
            formatted = self._format_html(headers, rows)
        elif fmt == "markdown":
            formatted = self._format_markdown(headers, rows)
        elif fmt == "csv":
            formatted = self._format_csv(headers, rows)
        else:
            formatted = self._format_html(headers, rows)

        return FormatTableOutput(formatted=formatted, row_count=len(rows))

    def _format_html(self, headers: list[str], rows: list[list[Any]]) -> str:
        """Format as HTML table."""
        html = ["<table class='data-table'>"]

        # Header
        html.append("<thead><tr>")
        for h in headers:
            html.append(f"<th>{h}</th>")
        html.append("</tr></thead>")

        # Body
        html.append("<tbody>")
        for row in rows:
            html.append("<tr>")
            for cell in row:
                html.append(f"<td>{cell}</td>")
            html.append("</tr>")
        html.append("</tbody>")

        html.append("</table>")
        return "\n".join(html)

    def _format_markdown(self, headers: list[str], rows: list[list[Any]]) -> str:
        """Format as Markdown table."""
        lines = []

        # Header
        lines.append("| " + " | ".join(headers) + " |")

        # Separator
        lines.append("| " + " | ".join(["---"] * len(headers)) + " |")

        # Rows
        for row in rows:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")

        return "\n".join(lines)

    def _format_csv(self, headers: list[str], rows: list[list[Any]]) -> str:
        """Format as CSV."""
        lines = [",".join(f'"{h}"' for h in headers)]

        for row in rows:
            lines.append(",".join(f'"{cell}"' for cell in row))

        return "\n".join(lines)


class AssembleDocumentTool(BaseTool):
    """Assemble sections into a complete document."""

    name = "assemble_document"
    category = ToolCategory.GENERATION
    description = "Assembles business case sections into PDF, DOCX, or HTML"
    input_schema = AssembleDocumentInput
    output_schema = AssembleDocumentOutput
    timeout_seconds = 30

    async def execute(self, input_data: AssembleDocumentInput) -> AssembleDocumentOutput:
        """Assemble complete document."""
        sections = input_data.sections
        output_format = input_data.output_format
        branding = input_data.branding or {}

        # Calculate metrics
        page_count = len(sections) + 2  # Title page + sections

        if output_format == "pdf":
            doc_bytes = await self._generate_pdf(sections, branding)
        elif output_format == "docx":
            doc_bytes = await self._generate_docx(sections, branding)
        elif output_format == "html":
            doc_bytes = await self._generate_html(sections, branding)
        else:
            doc_bytes = await self._generate_pdf(sections, branding)

        file_size = len(doc_bytes) if doc_bytes else 0

        return AssembleDocumentOutput(
            document_bytes=doc_bytes,
            document_url=None,  # Would upload to storage in production
            page_count=page_count,
            file_size_bytes=file_size,
        )

    async def _generate_pdf(self, sections: list[dict], branding: dict) -> bytes:
        """Generate PDF using WeasyPrint (or return HTML if unavailable)."""
        if not WEASYPRINT_AVAILABLE:
            logger.warning("weasyprint not available, returning HTML instead of PDF")
            html_content = await self._generate_html_content(sections, branding)
            return html_content.encode("utf-8")

        html_content = await self._generate_html_content(sections, branding)
        html = HTML(string=html_content)
        pdf_bytes = html.write_pdf()
        return pdf_bytes

    async def _generate_docx(self, sections: list[dict], branding: dict) -> bytes:
        """Generate DOCX using python-docx."""
        from docx import Document

        doc = Document()

        # Title
        doc.add_heading("Business Case", 0)

        # Add sections
        for section in sections:
            heading = section.get("title", "Section")
            content = section.get("content", "")

            doc.add_heading(heading, level=1)
            doc.add_paragraph(content)

        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    async def _generate_html(self, sections: list[dict], branding: dict) -> bytes:
        """Generate HTML document."""
        html = await self._generate_html_content(sections, branding)
        return html.encode("utf-8")

    async def _generate_html_content(self, sections: list[dict], branding: dict) -> str:
        """Generate HTML content."""
        company_name = branding.get("company_name", "Company Name")

        html_parts = [
            "<!DOCTYPE html>",
            "<html>",
            "<head>",
            "<title>Business Case</title>",
            "<style>",
            "body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }",
            "h1 { color: #2c3e50; border-bottom: 2px solid #3498db; }",
            "h2 { color: #34495e; margin-top: 30px; }",
            ".cover { text-align: center; padding: 100px 0; }",
            ".cover h1 { font-size: 36px; border: none; }",
            "</style>",
            "</head>",
            "<body>",
        ]

        # Cover page
        html_parts.extend(
            [
                "<div class='cover'>",
                "<h1>Business Case</h1>",
                f"<p>Prepared for {company_name}</p>",
                f"<p>{branding.get('date', '2024')}</p>",
                "</div>",
            ]
        )

        # Sections
        for section in sections:
            title = section.get("title", "Section")
            content = section.get("content", "")

            html_parts.append(f"<h2>{title}</h2>")
            html_parts.append(f"<div>{content}</div>")

        html_parts.extend(["</body>", "</html>"])

        return "\n".join(html_parts)
